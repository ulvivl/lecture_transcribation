import json
import re
from contextlib import contextmanager
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

SENTENCE1 = "Тема лекции"
SENTENCE2 = "Вопросы, которые будут рассматриваться в лекции"
SENTENCE3 = "Актуальность и значимость темы"


def dump_json(file: Path, data: Any) -> None:
    with open(file, "w") as fp:
        json.dump(data, fp=fp, ensure_ascii=False)


def load_json(file: Path) -> Any:
    with open(file) as fp:
        return json.load(fp=fp)


@contextmanager
def FileLock(lock_file: Path):
    lock_file.touch(exist_ok=False)
    try:
        yield
    finally:
        lock_file.unlink()


def format_glosary(glossary: dict, timecodes: dict = None) -> Document:
    document = Document()
    document.add_heading("Ключевые термины из лекции", 0)

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style("CommentsStyle", WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(16)
    obj_font.name = "Times New Roman"

    obj_charstyle_def = obj_styles.add_style("CommentsStyleDef", WD_STYLE_TYPE.CHARACTER)
    obj_font_def = obj_charstyle_def.font
    obj_font_def.size = Pt(15)
    obj_font_def.name = "Times New Roman"

    obj_charstyle_time = obj_styles.add_style("CommentsStyleTime", WD_STYLE_TYPE.CHARACTER)
    obj_font_time = obj_charstyle_time.font
    obj_font_time.size = Pt(13)
    obj_font_time.name = "Times New Roman"

    for term, definition in glossary.items():
        clear_def = definition.find("-")
        if clear_def != -1:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(term, style="CommentsStyle").bold = True
            paragraph.add_run(" -", style="CommentsStyleDef")
            paragraph.add_run(re.sub("-", "", definition[clear_def + 1 :]), style="CommentsStyleDef")
            if timecodes is not None:
                paragraph.add_run(f".\n\tВстречается в лекции: {timecodes[term]}\n", style="CommentsStyleTime")

    return document


def format_test(document: Document, tests: dict) -> Document:
    document = Document()
    document.add_heading("Тест по материалу из лекции", 0)

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style("CommentsStyle", WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(13)
    obj_font.name = "Times New Roman"

    for quest, ans in tests.items():
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run("Вопрос:\t", style="CommentsStyle").bold = True
        paragraph.add_run(f"{quest}\n", style="CommentsStyle")
        paragraph.add_run("Ответ:\t", style="CommentsStyle").bold = True
        paragraph.add_run(f"{ans}\n", style="CommentsStyle")

    return document


def format_docx(introduction, main_body, conclusion, tests):
    document = Document()
    document.add_heading("Конспект лекции", 0)
    document.add_heading(f"Введение", level=1)

    ind_start1 = introduction.find(SENTENCE1)
    if ind_start1 == -1:
        document.add_paragraph(introduction)
        document.add_heading(f"Основная часть", level=1)
        document.add_paragraph(main_body)

        document.add_heading(f"Заключение", level=1)
        document.add_paragraph(conclusion)
        return format_test(document=document, tests=tests) if tests else document

    introduction = introduction[ind_start1:]
    document.add_heading(f"{SENTENCE1}:", level=3)
    introduction = re.sub(f"{SENTENCE1}:", "", introduction)

    ind_start2 = introduction.find(SENTENCE2)
    if ind_start2 == -1:
        document = Document()
        document.add_heading("Конспект лекции", 0)
        document.add_heading(f"Введение", level=1)
        document.add_paragraph(introduction)
        document.add_heading(f"Основная часть", level=1)
        document.add_paragraph(main_body)

        document.add_heading(f"Заключение", level=1)
        document.add_paragraph(conclusion)
        return format_test(document=document, tests=tests) if tests else document
    document.add_paragraph(introduction[:ind_start2].strip())

    introduction = introduction[ind_start2:]
    introduction = re.sub(f"{SENTENCE2}:", "", introduction)

    document.add_heading(f"{SENTENCE2}:", level=3)
    ind_start3 = introduction.find(SENTENCE3)
    if ind_start2 == -1:
        document = Document()
        document.add_heading("Конспект лекции", 0)
        document.add_heading(f"Введение", level=1)
        document.add_paragraph(introduction)
        document.add_heading(f"Основная часть", level=1)
        document.add_paragraph(main_body)

        document.add_heading(f"Заключение", level=1)
        document.add_paragraph(conclusion)
        return format_test(document=document, tests=tests) if tests else document

    document.add_paragraph(introduction[:ind_start3].strip())

    introduction = introduction[ind_start3:]
    introduction = re.sub(f"{SENTENCE3}:", "", introduction)
    document.add_heading(f"{SENTENCE3}:", level=3)

    document.add_paragraph(introduction[:].strip())
    # document.add_page_break()

    # Обработка основной части
    themes, paragraphs = main_body
    document.add_heading(f"Основная часть", level=1)
    for theme, paragraph in zip(themes, paragraphs):
        document.add_heading(theme, level=3)
        document.add_paragraph("\t" + paragraph)

    document.add_heading(f"Заключение", level=1)
    document.add_paragraph("\t" + conclusion)

    return format_test(document=document, tests=tests) if tests else document
